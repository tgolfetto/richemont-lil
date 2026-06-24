from docx import Document
from docx.shared import Pt


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    p.paragraph_format.space_after = Pt(10)


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)


def main() -> None:
    doc = Document()

    h1(doc, "Richemont Learning Platform\nData Structure Document")
    body(doc, "Level: moderately technical (business-readable)")

    h2(doc, "1. Data Model Intent")
    body(
        doc,
        "The model is designed to support campaign targeting, recommendation relevance, and measurable learning progression across multiple dimensions: role, market, industry, language, employee level, and skill proficiency."
    )

    h2(doc, "2. Core Entities")
    bullet(doc, "Users: learning managers and employees with role, market, job title, and optional spoken languages.")
    bullet(doc, "Campaigns: learning initiatives with targeting properties, timeline, status, and structured learning preferences.")
    bullet(doc, "Skills: capability topics used to define learning focus.")
    bullet(doc, "Courses: learning assets linked to a skill, with level and content metadata.")
    bullet(doc, "Recommendations: assignments connecting users to courses with lifecycle status.")

    h2(doc, "3. Campaign Structure")
    body(doc, "Campaigns hold both operational and targeting properties as structured fields.")
    bullet(doc, "Operational: name, description, status, start_date, end_date, created_by.")
    bullet(doc, "Audience targeting: target_role, target_market, target_levels.")
    bullet(doc, "Learning targeting: campaign_type, focus_skills, skill_proficiency_levels, skill_matrix.")
    bullet(doc, "Localization targeting: target_languages.")
    body(
        doc,
        "This avoids relying on free-form description text and enables consistent filtering, editing, and exports."
    )

    h2(doc, "4. Recommendation Lifecycle")
    body(doc, "Recommendations track learner progress through a simple state machine:")
    bullet(doc, "Assigned")
    bullet(doc, "In Progress")
    bullet(doc, "Completed")
    body(
        doc,
        "This lifecycle is used in manager dashboards, campaign details, and employee views. It also powers status exports and popup user lists."
    )

    h2(doc, "5. Skill and Proficiency Design")
    body(
        doc,
        "Each campaign can define multiple skills and an expected proficiency for each skill. This is a key modeling choice to keep recommendations aligned with development intent."
    )
    bullet(doc, "Skill matrix format conceptually represents pairs: skill -> proficiency.")
    bullet(doc, "Proficiency levels include: General, Beginner, Beginner + Intermediate, Intermediate, Advanced.")
    bullet(doc, "Course selection is filtered by proficiency rules before campaign-type logic is applied.")

    h2(doc, "6. Campaign Type Logic")
    bullet(doc, "Tailored: maximum 10 courses per skill x language.")
    bullet(doc, "Exhaustive: all available courses per skill x language.")
    body(
        doc,
        "This creates a predictable distinction between curated campaigns and full-catalog campaigns."
    )

    h2(doc, "7. Language Handling")
    body(
        doc,
        "Language is selected at campaign level, not per individual recommendation. The system then presents course recommendations by selected language context."
    )
    bullet(doc, "Supports multilingual rollout without duplicating campaign setup.")
    bullet(doc, "Allows course display and exports to remain consistent with campaign language scope.")

    h2(doc, "8. Employee-Side Data Consumption")
    body(doc, "Employee pages consume recommendation rows enriched with context:")
    bullet(doc, "Course details: title, skill, level, watch link.")
    bullet(doc, "Campaign source: campaign_name.")
    bullet(doc, "Language context: language used in presentation.")
    bullet(doc, "Status separation: in-progress, assigned, completed.")

    h2(doc, "9. Manager-Side Data Consumption")
    body(doc, "Manager flows rely on structured campaign and recommendation data for:")
    bullet(doc, "Campaign detail analytics (A/P/C counts).")
    bullet(doc, "Course-by-skill views with status distribution.")
    bullet(doc, "Employee lists per course status in popup tables.")
    bullet(doc, "CSV exports (campaign details and filtered user status lists).")

    h2(doc, "10. Compatibility and Migration Strategy")
    body(
        doc,
        "The platform currently includes backward-compatible behavior: if newly introduced campaign columns are not yet present in Supabase, APIs can fall back to base field writes."
    )
    bullet(doc, "Preferred path: migrate schema so structured fields are available.")
    bullet(doc, "Fallback path: keep core campaign operations functional during transition.")

    h2(doc, "11. Summary")
    body(
        doc,
        "The data structure intentionally balances governance and usability. Managers get structured control and reporting, while employees get focused, relevant learning flows with clear progress states."
    )

    doc.save("Richemont_Data_Structure_Document.docx")


if __name__ == "__main__":
    main()
