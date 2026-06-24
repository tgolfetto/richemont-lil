from docx import Document
from docx.shared import Pt


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    p.space_before = Pt(12)
    p.space_after = Pt(6)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)


def main() -> None:
    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("Richemont Learning Platform\nData Structure Overview")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.add_paragraph("Business-oriented summary of design choices")
    subtitle.runs[0].italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    add_heading(doc, "Purpose")
    add_body(
        doc,
        "The platform data structure was designed to support practical learning operations at scale while keeping the experience flexible for both Learning Managers and employees. "
        "The goal is to make campaign design, course targeting, and learner progress easy to manage across roles, markets, industries, and languages."
    )

    add_heading(doc, "Key Design Choices")
    add_bullet(
        doc,
        "Skills are tied to proficiency expectations. Each campaign can define which skills matter and the desired proficiency level for each one."
    )
    add_bullet(
        doc,
        "Language is set at campaign level. This allows one campaign to serve multilingual audiences consistently and generate recommendations per selected language."
    )
    add_bullet(
        doc,
        "Campaign type is explicit: Tailored campaigns focus on a curated set (up to 10 courses per skill and language), while Exhaustive campaigns expose the full available catalog."
    )
    add_bullet(
        doc,
        "Targeting is multi-dimensional. Campaigns can be aimed by role hierarchy, market hierarchy, industry type, employee level, language, and skill focus."
    )
    add_bullet(
        doc,
        "Employee progress is tracked as a clear journey: Assigned, In Progress, Completed. This structure supports both operational monitoring and simple communication."
    )

    add_heading(doc, "Learning Manager Perspective")
    add_body(
        doc,
        "From the manager side, the structure supports campaign creation and editing with meaningful controls rather than generic fields. "
        "Managers can shape campaigns by audience and expected capability uplift, then monitor outcomes by campaign, market, and course status."
    )
    add_bullet(
        doc,
        "Campaign properties are treated as first-class data (not just text notes), improving consistency over time."
    )
    add_bullet(
        doc,
        "Course-level visibility includes employee status lists, enabling direct follow-up and accountability."
    )
    add_bullet(
        doc,
        "Exports support practical reporting needs for campaign details and user status breakdowns."
    )

    add_heading(doc, "Employee Perspective")
    add_body(
        doc,
        "On the employee side, the model is intentionally simple: a clear Courses area and a dedicated Learning Path area. "
        "Employees can focus on what is assigned now, what is in progress, and what is completed, while also selecting skills they want to improve."
    )
    add_bullet(
        doc,
        "Learning Path recommendations adapt to selected skills."
    )
    add_bullet(
        doc,
        "Skill levels and milestones make progression visible and action-oriented."
    )
    add_bullet(
        doc,
        "Course context includes level, language, and campaign source so employees understand why each course appears."
    )

    add_heading(doc, "Why This Structure Works")
    add_body(
        doc,
        "This approach balances strategic control and day-to-day usability. It gives Learning Managers the structure needed for governance and reporting, while giving employees a focused, intuitive path for development. "
        "Most importantly, it keeps learning recommendations aligned with real business context: role, market, language, and skill maturity."
    )

    output = "Richemont_Data_Structure_Overview.docx"
    doc.save(output)


if __name__ == "__main__":
    main()
