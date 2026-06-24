from docx import Document
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    p.paragraph_format.space_after = Pt(10)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.paragraph_format.space_after = Pt(12)


def add_section(doc: Document, heading: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.bold = True
    run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)


def main() -> None:
    doc = Document()

    add_title(doc, "Product Requirements Document (PRD)")
    add_subtitle(doc, "Richemont Learning & Development Platform")

    add_section(doc, "1. Product Vision")
    add_body(
        doc,
        "Create a unified learning platform that enables Learning Managers to design high-impact campaigns and helps employees follow a clear, personalized development journey across markets, roles, and languages."
    )

    add_section(doc, "2. Problem Statement")
    add_body(
        doc,
        "Learning campaigns are often difficult to tailor at scale across regions and employee profiles. Employees may receive recommendations that are generic, not aligned to their language or proficiency level, and disconnected from clear progression milestones."
    )

    add_section(doc, "3. Goals")
    add_bullet(doc, "Enable managers to create, edit, monitor, export, and delete campaigns with confidence.")
    add_bullet(doc, "Ensure campaigns are targeted by role, market, industry, language, employee level, and skill proficiency.")
    add_bullet(doc, "Provide employees with a clear courses view (assigned, in progress, completed) and a dedicated learning-path experience.")
    add_bullet(doc, "Increase recommendation relevance and measurable completion outcomes.")

    add_section(doc, "4. Primary Users")
    add_body(doc, "Learning Manager")
    add_bullet(doc, "Designs campaign targeting and expected skill uplift.")
    add_bullet(doc, "Monitors campaign progress and learner status.")
    add_bullet(doc, "Exports campaign and user-level data for reporting.")
    add_body(doc, "Employee")
    add_bullet(doc, "Consumes assigned and in-progress courses.")
    add_bullet(doc, "Selects desired improvement skills in Learning Path.")
    add_bullet(doc, "Tracks milestones and re-engages completed courses.")

    add_section(doc, "5. Scope")
    add_body(doc, "In Scope")
    add_bullet(doc, "Campaign creation/editing with structured targeting properties.")
    add_bullet(doc, "Campaign details with skill-based recommended course lists.")
    add_bullet(doc, "Tailored vs Exhaustive campaign logic.")
    add_bullet(doc, "Employee Courses page and separate Learning Path page.")
    add_bullet(doc, "Exports for campaign details and user status lists.")
    add_body(doc, "Out of Scope (Current Phase)")
    add_bullet(doc, "Automated LMS integrations beyond LinkedIn Learning links.")
    add_bullet(doc, "Advanced AI-based adaptive sequencing in real time.")
    add_bullet(doc, "Formal certification workflows.")

    add_section(doc, "6. Core Product Choices")
    add_bullet(doc, "Each campaign can define a set of target skills and a proficiency level per skill.")
    add_bullet(doc, "Campaign languages are selected at campaign level and used to shape course presentation.")
    add_bullet(doc, "Tailored campaigns show up to 10 courses per skill per language.")
    add_bullet(doc, "Exhaustive campaigns show all available courses per skill per language.")
    add_bullet(doc, "Employee course status follows a simple lifecycle: Assigned, In Progress, Completed.")
    add_bullet(doc, "Employee-facing navigation separates Courses and Learning Path to reduce cognitive load.")

    add_section(doc, "7. Functional Requirements")
    add_body(doc, "Campaign Management")
    add_bullet(doc, "Managers can create campaigns with structured properties and timeline.")
    add_bullet(doc, "Managers can edit campaign properties, including skill-proficiency pairs.")
    add_bullet(doc, "Managers can delete campaigns from edit view.")
    add_bullet(doc, "Campaign details show course-level status distribution (A / P / C) and employee lists.")
    add_body(doc, "Recommendation Logic")
    add_bullet(doc, "Recommendations are grouped by skill and language.")
    add_bullet(doc, "Course lists are paginated to keep detail pages compact.")
    add_bullet(doc, "Employees linked to each course status are viewable in popup with export option.")
    add_body(doc, "Employee Experience")
    add_bullet(doc, "Courses page separates In Progress and Assigned sections, with larger focus on In Progress.")
    add_bullet(doc, "Completed courses include a Rewatch action.")
    add_bullet(doc, "Learning Path allows skill selection and displays suggested courses based on selected interests.")
    add_bullet(doc, "Course cards show key context: level, language, and source campaign.")

    add_section(doc, "8. Analytics and Success Metrics")
    add_bullet(doc, "Activation Rate")
    add_bullet(doc, "Recommendation Acceptance Rate")
    add_bullet(doc, "Course Completion Rate")
    add_bullet(doc, "Monthly Active Learners")
    add_bullet(doc, "Campaign Completion by Market")
    add_bullet(doc, "Skill Impact progression indicators")

    add_section(doc, "9. UX Principles")
    add_bullet(doc, "Use a consistent flat visual system across manager and employee areas.")
    add_bullet(doc, "Keep primary actions visible and predictable.")
    add_bullet(doc, "Prefer structured data fields over free-form metadata text.")
    add_bullet(doc, "Present information density where needed, but keep navigation simple.")

    add_section(doc, "10. Release Readiness Criteria")
    add_bullet(doc, "Managers can complete end-to-end campaign lifecycle (create, edit, export, delete).")
    add_bullet(doc, "Employees can navigate Courses and Learning Path without ambiguity.")
    add_bullet(doc, "Campaign type logic is respected in course list generation.")
    add_bullet(doc, "Exports contain expected campaign and learner status information.")
    add_bullet(doc, "Core flows are stable in local and connected Supabase mode.")

    add_section(doc, "11. Risks and Mitigations")
    add_bullet(doc, "Schema drift risk: use backward-compatible API behavior and migration scripts.")
    add_bullet(doc, "Data quality risk in recommendations: rely on structured campaign properties.")
    add_bullet(doc, "UI complexity risk: separate employee experiences into focused pages.")

    output = "Richemont_PRD.docx"
    doc.save(output)


if __name__ == "__main__":
    main()
